from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, Inches
from docx.enum.section import WD_ORIENT

from thematic import Thematic
from complex_theme import ComplexTheme
import complex_theme_private_collection
import temporary_team_private_collection
from worksheet import Worksheet
from worksheet_line import WorksheetLine
import worksheet_line_private_collection as wsl
import outlay_tree

#===================================================================================================

# Common functions.

def set_table_columns_widths(t, ws):
    """
    Set table columns widths.

    Parameters
    ----------
    t : Table
        Word table.
    ws : [float]
        Columns widths in inches.
    """

    ws = [Inches(w) for w in ws]

    for row in t.rows:
        for i, w in enumerate(ws):
            row.cells[i].width = w

#===================================================================================================

class GeneratorWord:
    """
    Master for word documents generation.
    """

    """
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')
    """

    #-----------------------------------------------------------------------------------------------
    # Common methods.
    #-----------------------------------------------------------------------------------------------

    def __init__(self):
        """
        Init document.
        """

        self.doc = Document()
        sections = self.doc.sections

        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.25)
            section.right_margin = Cm(1.25)

    #-----------------------------------------------------------------------------------------------

    def save(self, name):
        """
        Save document.

        Parameters
        ----------
        name : str
            File name.
        """

        self.doc.save(name)

    #-----------------------------------------------------------------------------------------------

    def add_empty_line(self):
        """
        Add empty line.
        """

        self.doc.add_paragraph('')

    #-----------------------------------------------------------------------------------------------

    def add_empty_lines(self, n):
        """
        Add empty lines.

        Parameters
        ----------
        n : int
            Count.
        """

        for _ in range(n):
            self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_paragraph(self, text, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, is_bold=False):
        """
        Add paragraph.

        Parameters
        ----------
        text : str
            Text.
        alignment : any
            Text alignment.
        is_bold : bool
            Bold check.

        Returns
        -------
        Paragraph
            Paragraph.
        """

        p = self.doc.add_paragraph()
        r = p.add_run(text)
        p.alignment = alignment
        r.bold = is_bold
        f = p.style.font
        f.name = 'Times New Roman'
        f.size = Pt(12)

        return p

    #-----------------------------------------------------------------------------------------------

    def add_corner_inscription_supplement_to_order(self, n):
        """
        Add inscription to corner of document.
        Inscription contains text about supplement of order.

        Parameters
        ----------
        n : int
            Number of supplement.
        """

        p = self.doc.add_paragraph(f'Приложение № {n}\n'
                                   'к приказу НИЦ «Курчатовский институт»\n'
                                   'от «___» ____________ ____ г. № ______')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        f = p.style.font
        f.name = 'Times New Roman'
        f.size = Pt(12)

    #-----------------------------------------------------------------------------------------------

    def add_signature(self, w, add_line_for_signature):
        """
        Add signature.

        Parameters
        ----------
        w : WorksheetLine
            Worksheet line.
        add_line_for_signature : bool
            Add line for signature.
        """

        # Table and its style.
        t = self.doc.add_table(rows=1, cols=5)
        h = t.rows[0].cells

        # Take job title with first big letter
        title_name = w.job_title.name
        if (title_name == 'руководитель отделения') or (title_name == 'директор департамента'):
            title_name = title_name.split()[0]
        title_name = title_name[0:1].upper() + title_name[1:]
        h[0].text = f'{title_name} {w.job_place.name_r}'

        h[1].text = ''
        if add_line_for_signature:
            h[2].text = '_____________________\n'\
                        '      (подпись)'
        h[3].text = ''
        h[4].text = w.employee.personal.surname_n_p()

        # Cells sizes.
        xs = [Inches(3.0), Inches(0.5), Inches(1.0), Inches(0.5), Inches(1.5)]
        for row in t.rows:
            for i, x in enumerate(xs):
                row.cells[i].width = x

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_signatures(self, ws, add_line_for_signature=True):
        """
        Add signatures.

        Parameters
        ----------
        ws : [WorksheetLine]
            List of worksheet lines.
        add_line_for_signature : bool
            Add line for signature.
        """

        for w in ws:
            self.add_signature(w, add_line_for_signature)

    #-----------------------------------------------------------------------------------------------
    # Technical task methods.
    #-----------------------------------------------------------------------------------------------

    def add_technical_task_title(self, cx, y):
        """
        Add title for technical task.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Start year.
        """

        text = 'ТЕХНИЧЕСКОЕ ЗАДАНИЕ\n'\
               'на выполнение научно-исследовательской работы\n'\
               f'по комплексной теме {cx.title}'
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #-----------------------------------------------------------------------------------------------

    def add_thematic_results_table_with_TRL(self, th, y):
        """
        Add thematic results table with TRL.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year.
        """

        # Calculate rows count.
        rows_count = 0
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                rows_count = rows_count + 1
        if rows_count == 0:
            return

        # Table and its style.
        t = self.doc.add_table(rows=rows_count + 1, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Планируемый результат'
        h[2].text = 'Планируемый УГТ'
        h[3].text = 'Год'

        # Write all results.
        i = 0
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                h = t.rows[1 + i].cells
                h[0].text = str(1 + i)
                h[1].text = r.title
                h[2].text = '1'
                h[3].text = str(r.year)
                i = i + 1

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_rids_table_with_TRL(self, th, y):
        """
        Add thematic rids table with TRL.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year
        """

        # Calculate rows count.
        rows_count = 0
        for r in th.results:
            if r.is_rid and (r.year in range(y, y + 3)):
                rows_count = rows_count + 1
        if rows_count == 0:
            return

        # Table and its style.
        t = self.doc.add_table(rows=rows_count + 1, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Планируемый результат'
        h[2].text = 'Планируемый УГТ'
        h[3].text = 'Год'

        # Write all results.
        i = 0
        for r in th.results:
            if r.is_rid and (r.year in range(y, y + 3)):
                h = t.rows[1 + i].cells
                h[0].text = str(1 + i)
                h[1].text = f'Планируемый вид РИД: {r.rid_type}.\n'\
                            f'Планируемое название РИД: «{r.rid_name}\n'\
                            f'Планируемый результат: {r.description}'
                h[2].text = '3'
                h[3].text = str(r.year)
                i = i + 1

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_indicators_table(self, th, y):
        """
        Add thematic indicators table.

        Parameters
        ----------
        th : Thematic
            Thematic.
        y : int
            Year.
        """

        # Table and its style.
        t = self.doc.add_table(rows=5, cols=6)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Показатель (индикатор)'
        h[2].text = 'Единица измерения'
        h[3].text = f'Год {y}'
        h[4].text = f'Год {y + 1}'
        h[5].text = f'Год {y + 2}'

        # Doctors.
        h = t.rows[1].cells
        h[0].text = '1'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени доктора наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_doctors(y + i)}'

        # Candidates.
        h = t.rows[2].cells
        h[0].text = '2'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени кандидата наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_candidates(y + i)}'

        # Rids.
        h = t.rows[3].cells
        h[0].text = '3'
        h[1].text = 'Количество полученных результатов интеллектуальной деятельности'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_rids (y + i)}'

        # Publications.
        h = t.rows[4].cells
        h[0].text = '4'
        h[1].text = 'Публикации в журналах, индексируемых в российских и международных '\
                    'информационно-аналитических системах научного цитирования (Российский индекс '\
                    'научного цитирования или публикации в научных журналах, '\
                    'входящих в «Белый список»'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{th.ind_publications(y + i)}'

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5, 0.5, 0.5])

        self.add_empty_line()

    #-----------------------------------------------------------------------------------------------

    def add_thematic_characteristics(self, th, number, y):
        """
        Add thematic characteristics.

        Parameters
        ----------
        th : Thematic
            Thematic.
        number : int
            Number in enumeration.
        y : int
            Year.
        """

        self.add_paragraph(f'7.{number}. Тематика исследований {th.title}.')

        # p. 1-4.
        self.add_paragraph(f'1) Цель работы: {th.goal}')
        self.add_paragraph(f'2) Актуальность и новизна работы: {th.actuality}')
        self.add_paragraph(f'3) Ресурсная обеспеченность работы: {th.resources}')
        self.add_paragraph(f'4) Имеющийся научно-технический задел по работе: {th.background}')

        # p. 5. content
        cs = []
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                cs.append(r.content)
        txt = ' '.join(cs)
        self.add_paragraph(f'5) Основное содержание работ: {txt}')

        # p. 6. years
        self.add_paragraph(f'6) Срок выполнения работы: {y}-{y + 2} годы.')

        # p. 7. results.
        rs = []
        for r in th.results:
            if (not r.is_rid) and (r.year in range(y, y + 3)):
                rs.append(r.title)
        txt = ' '.join(rs)
        self.add_paragraph(f'7) Планируемые результаты: {txt}')
        self.add_thematic_results_table_with_TRL(th, y)
        self.add_thematic_rids_table_with_TRL(th, y)

        # p. 8. indicators
        self.add_paragraph('8) Значение целевых индикаторов и показателей')
        self.add_thematic_indicators_table(th, y)

    #-----------------------------------------------------------------------------------------------

    def add_complex_theme_indicators_table(self, cx, y):
        """
        Add complex theme indicators table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # Name.
        self.add_paragraph('6. Значение показателей, характеризующих качество работ:')

        # Table and its style.
        t = self.doc.add_table(rows=5, cols=6)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Показатель (индикатор)'
        h[2].text = 'Единица измерения'
        h[3].text = f'Год {y}'
        h[4].text = f'Год {y + 1}'
        h[5].text = f'Год {y + 2}'

        # Doctors.
        h = t.rows[1].cells
        h[0].text = '1'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени доктора наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_doctors(y + i)}'

        # Candidates.
        h = t.rows[2].cells
        h[0].text = '2'
        h[1].text = 'Количество защищенных диссертаций на соискание ученой степени кандидата наук'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_candidates(y + i)}'

        # Rids.
        h = t.rows[3].cells
        h[0].text = '3'
        h[1].text = 'Количество полученных результатов интеллектуальной деятельности'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_rids (y + i)}'

        # Publications.
        h = t.rows[4].cells
        h[0].text = '4'
        h[1].text = 'Публикации в журналах, индексируемых в российских и международных '\
                    'информационно-аналитических системах научного цитирования (Российский индекс '\
                    'научного цитирования или публикации в научных журналах, '\
                    'входящих в «Белый список»'
        h[2].text = 'шт'
        for i in range(3):
            h[3 + i].text = f'{cx.ind_publications(y + i)}'

        set_table_columns_widths(t, [0.5, 5.0, 0.5, 0.5, 0.5, 0.5])

    #-----------------------------------------------------------------------------------------------

    def add_complex_theme_characteristics(self, cx, y):
        """
        Add complex theme characteristics.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # p. 1-4.
        self.add_paragraph('1. Подтемы комплексной темы:')
        self.add_paragraph(f'2. Цель выполнения НИР: {cx.goal}')
        self.add_paragraph(f'3. Срок выполнения НИР: {y}-{y + 2} годы.')
        self.add_paragraph(f'4. Актуальность НИР: {cx.actuality}')

        # p. 5. Get all results (without RIDs).
        rs = []
        for th in cx.thematics:
            for r in th.results:
                if (not r.is_rid) and (r.year in range(y, y + 3)):
                    rs.append(r.title)
        txt = ' '.join(rs)
        self.add_paragraph(f'5. Планируемые результаты НИР: {txt}')

        # p. 6.
        self.add_complex_theme_indicators_table(cx, y)

        # p. 6.
        self.add_empty_line()
        self.add_paragraph('7. Основное содержание работ по тематикам исследований:')

        # Add characteristics for all thematics.
        for i, th in enumerate(cx.thematics):
            self.add_thematic_characteristics(th, i + 1, y)

    #-----------------------------------------------------------------------------------------------
    # Calendar plan methods.
    #-----------------------------------------------------------------------------------------------

    def add_calendar_plan_title(self, cx, y):
        """
        Add title for calendar_plan.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Start year.
        """

        text = 'КАЛЕНДАРНЫЙ ПЛАН\n'\
               f'на {y} год и плановый период {y + 1} и {y + 2} годов на выполнение '\
               f'научно-исследовательской работы по комплексной теме {cx.title}'
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #-----------------------------------------------------------------------------------------------

    def add_calendar_plan_table(self, cx, y):
        """
        Add calendar plan table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        # Calculate all results in year.
        k = len(cx.thematics)
        for th in cx.thematics:
            for r in th.results:
                if r.year == y:
                    k = k + 1

        self.add_paragraph(f'Календарный план на {y} год', WD_PARAGRAPH_ALIGNMENT.CENTER, True)

        # Table and its style.
        t = self.doc.add_table(rows=1+k, cols=11)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Содержание выполняемых работ'
        h[2].text = 'Ожидаемый результат выполнения работ / Период реализации работ в рамках тематики исследований'
        h[3].text = 'Вид планируемого к созданию результата интеллектуальной деятельности'
        h[4].text = 'Планируемое наименование планируемого к созданию РИД'
        h[5].text = 'Краткое описание планируемого к созданию РИД'
        h[6].text = 'Планируемый уровень готовности разрабатываемых или разработанных технологий (далее - УГТ) **'
        h[7].text = 'Срок создания РИД (мм.гггг)'
        h[8].text = 'Состав отчетной документации'
        h[9].text = 'Ответственный руководитель работ по комплексной теме\n\n'\
                    'Ответственное структурное подразделение Центра за выполнение работ по комплексной теме\n\n'\
                    'Ответственный руководитель работ по подтеме комплексной темы с указанием структурного подразделения\n\n'\
                    'Ответственный руководитель работ по тематике исследований подтемы комплексной темы с указанием структурного подразделения'
        h[10].text = 'Стоимость, тыс. рублей'

        # Add row by row.
        i = 1
        for thi, th in enumerate(cx.thematics):
            h = t.rows[i].cells
            h[0].text = f'{thi + 1}.'
            h[1].merge(h[2]).merge(h[3]).merge(h[4]).merge(h[5]).merge(h[6]).merge(h[7]).merge(h[8]).merge(h[9])
            h[1].text = f'Тематика {th.title}'
            i = i + 1
            for r in th.results:
                if r.year == y:
                    h = t.rows[i].cells
                    h[0].text = ''
                    h[1].text = r.content
                    h[2].text = f'{r.title} / {y} год'
                    if r.is_rid:
                        h[3].text = r.rid_type
                        h[4].text = r.rid_name
                        h[5].text = r.description
                        h[6].text = '3'
                        h[7].text = f'12.{y}'
                    h[8].text = 'Аннотационный отчет - ежеквартально;\n\nИтоговый отчет о НИР.'
                    pname = cx.manager.employee.personal.surname_n_p('ru')
                    pjobtitle = cx.manager.job_title.name
                    pjobplace = cx.manager.job_place.name
                    h[9].text = f'{pname}\n({pjobtitle}, {pjobplace})'
                    money = round(cx.outlay.xmoney * (0.01 * r.funding_part), 2)
                    h[10].text = f'{money}'
                    i = i + 1

        # Text sizes.
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        font = run.font
                        font.size = Pt(8)

        #set_table_columns_widths(t, [0.5, 2.0, 2.0, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 2.0, 0.5])

        p = self.add_paragraph('* В настоящем столбце указываются следующие вида РИД - программы '
                               'для ЭВМ, базы данных, изобретения, полезные модели, промышленные '
                               'образцы, селекционные достижения, топологии интегральных микросхем, '
                               'секреты производства (ноу-хау), товарные знаки и знаки обслуживания, '
                               'коммерческие обозначения;', WD_PARAGRAPH_ALIGNMENT.LEFT)
        p.runs[0].font.size = Pt(8)
        p = self.add_paragraph('** Указывается планируемый результат в соответствии с Порядком '
                               'определения уровней готовности разрабатываемых или разработанных '
                               'технологий, а также научных и (или) научно-технических результатов, '
                               'соответствующих каждому уровню готовности технологий, утвержденным '
                               'приказом Минобрнауки России от 6 февраля 2023 г. № 107 и положениями '
                               'приказа НИЦ «Курчатовский институт» от 25 июля 2023 г. № 2136 '
                               '«Об организации в НИЦ «Курчатовский институт» работы по определению '
                               'уровней готовности разрабатываемых или разработанных технологи»:',
                               WD_PARAGRAPH_ALIGNMENT.LEFT)
        p.runs[0].font.size = Pt(8)

    #-----------------------------------------------------------------------------------------------
    # Outlay methods.
    #-----------------------------------------------------------------------------------------------

    def add_outlay_title(self, cx, y):
        """
        Add title for outlay.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Start year.
        """

        text = 'ПРЕДВАРИТЕЛЬНАЯ СМЕТА\n'\
               f'расходов на выполнение работы по комплексной теме {cx.title}\n'\
               f' на очередной {y} год и плановый период {y + 1} и {y + 2} годов'
        self.add_paragraph(text, WD_PARAGRAPH_ALIGNMENT.CENTER, True)

    #-----------------------------------------------------------------------------------------------

    def add_short_outlay_table(self, cx, y):
        """
        Add short outlay table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Year.
        """

        self.add_paragraph('тыс. рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        k = len(cx.thematics)

        # Table and its style.
        t = self.doc.add_table(rows=2*k + 4, cols=5)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование статей расходов'
        h[2].text = f'{y} год'
        h[3].text = f'{y + 1} год'
        h[4].text = f'{y + 2} год'

        # 1.
        h = t.rows[1].cells
        h[0].text = '1.'
        h[1].text = 'ВСЕГО по комплексной теме, в том числе'

        # 2.
        h = t.rows[2].cells
        h[0].text = '2.'
        h[1].text = 'Прямые и общепроизводственные затраты, из них:'

        # 3.
        h = t.rows[k + 3].cells
        h[0].text = '3.'
        h[1].text = 'Общехозяйственные расходы (*), из них:'

        # Process all thematics.
        direct, hoz = 0.0, 0.0
        for i, th in enumerate(cx.thematics):
            #
            h = t.rows[3 + i].cells
            h[0].text = f'2.{i + 1}.'
            h[1].text = f'Тематика исследований {th.title}'
            x = round(th.outlay['II'].xmoney, 2)
            direct = direct + x
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'
            #
            h = t.rows[4 + k + i].cells
            h[0].text = f'3.{i + 1}.'
            h[1].text = f'Тематика исследований {th.title}'
            x = round(th.outlay['III'].xmoney, 2)
            hoz = hoz + x
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'

        # Write sums.
        h = t.rows[2].cells
        direct = round(direct, 2)
        h[2].text = f'{direct}'
        h[3].text = f'{direct}'
        h[4].text = f'{direct}'
        h = t.rows[6].cells
        hoz = round(hoz, 2)
        h[2].text = f'{hoz}'
        h[3].text = f'{hoz}'
        h[4].text = f'{hoz}'
        h = t.rows[1].cells
        s = round(direct + hoz, 2)
        h[2].text = f'{s}'
        h[3].text = f'{s}'
        h[4].text = f'{s}'

        set_table_columns_widths(t, [0.5, 4.0, 1.0, 1.0, 1.0])

        self.add_paragraph('* Объем общехозяйственных расходов может быть изменен с учётом '
                           'изменений Плана доходов и расходов и принятых решений на '
                           'Финансовом совете', WD_PARAGRAPH_ALIGNMENT.LEFT)

    #-----------------------------------------------------------------------------------------------

    def add_outlay_table_theme_with_subthemes(self, outlay, y, hoz):
        """
        Add outlay table for one theme - subthemes.

        NB! We have no subthemes.

        Parameters
        ----------
        outlay : outlay_tree.Nod
            Outlay
        y : int
            Year.
        hoz : bool
            Add hoz spents.
        """

        self.add_paragraph('тыс. рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # Flatten outlay.
        outlay_lines = outlay.flatten()
        k = outlay['II'].count()

        # Table and its style.
        t = self.doc.add_table(rows=1+k, cols=3)
        t.style = 'Table Grid'

        # Add row.
        def add_row(i, h0, h1, h2):
            h = t.rows[i].cells
            h[0].text, h[1].text, h[2].text = h0, h1, h2

        # Add head.
        add_row(0,
                '№ п/п', 'Наименование статей расходов', 'Всего стоимость, тыс. рублей')

        # Add all lines.

        rowi = 1

        # If it is form without hoz then print first row.
        if not hoz:
            x = round(outlay['II'].xmoney, 2)
            add_row(rowi,'', 'ВСЕГО ЗАТРАТ, в том числе:', f'{x}')
            rowi = rowi + 1

        for ol in outlay_lines:

            # Do not print
            if (ol.label == 'I.') or (ol.label == 'II.'):
                continue

            # Do not print hoz spents.
            if (not hoz) and (ol.label == 'III.'):
                break

            x = round(ol.xmoney, 2)
            add_row(rowi, ol.label, ol.name, f'{x}')
            rowi = rowi + 1

        set_table_columns_widths(t, [0.5, 5.0, 1.5])

    #-----------------------------------------------------------------------------------------------

    def add_outlay_table(self, outlay, y):
        """
        Add outlay table.

        Parameters
        ----------
        outlay : outlay_tree.Nod
            Outlay.
        y : int
            Year.
        """

        self.add_paragraph('тыс. рублей', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        outlay_lines = outlay.flatten()
        k = len(outlay_lines)

        # Table and its style.
        t = self.doc.add_table(rows=1+k, cols=5)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование статей расходов'
        h[2].text = f'{y} год'
        h[3].text = f'{y + 1} год'
        h[4].text = f'{y + 2} год'

        # Add all lines.
        for i in range(k):
            ol = outlay_lines[i]
            h = t.rows[1 + i].cells
            h[0].text = ol.label
            h[1].text = ol.name
            x = round(ol.xmoney, 2)
            h[2].text = f'{x}'
            h[3].text = f'{x}'
            h[4].text = f'{x}'

        set_table_columns_widths(t, [0.5, 4.0, 1.0, 1.0, 1.0])

    #-----------------------------------------------------------------------------------------------
    # Temporary team.
    #-----------------------------------------------------------------------------------------------

    def add_temporary_team_title(self, cx, y):
        """
        Add title for temporary team.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        y : int
            Start year.
        """

        pre_text = 'СОСТАВ\n'\
                   'временного трудового коллектива\n'\
                   'для выполнения научно-исследовательской работы по комплексной теме\n'
        post_text = f' на очередной {y} год и плановый период {y + 1} и {y + 2} годов'
        p = self.doc.add_paragraph()
        r = p.add_run(pre_text + cx.title + post_text)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #-----------------------------------------------------------------------------------------------

    def add_temporary_team_table_with_workload(self, w, cx):
        """
        Add temporary team table with workload.

        Parameters
        ----------
        w : Worksheet
            Worksheet.
        cx : ComplexTheme
            Complex theme.
        """

        # number of people
        n = len(w.lines)

        self.add_paragraph('человеко/месяц', WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # table and its style
        t = self.doc.add_table(rows=n + 1, cols=9)
        t.style = 'Table Grid'

        # head
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'ФИО'
        h[2].text = 'Должность'
        h[3].text = 'Табельный номер'
        h[4].text = 'Год рождения'
        h[5].text = 'Подразделение'
        h[6].text = f'Планируемые трудозатраты по тематике {cx.thematics[0].title}'
        h[7].text = f'Планируемые трудозатраты по тематике {cx.thematics[1].title}'
        h[8].text = f'Планируемые трудозатраты по тематике {cx.thematics[2].title}'

        # add rest rows
        for i in range(n):
            r = t.rows[i + 1].cells
            wl = w.lines[i]
            e = wl.employee
            p = e.personal
            r[0].text = str(i + 1) + '.'
            r[1].text = p.surname_name_patronymic()
            r[2].text = wl.job_title.name
            r[3].text = e.tabel
            r[4].text = str(p.year)
            r[5].text = wl.job_place.half_full_name
            x = round(wl.slot / 3.0, 2)
            r[6].text = f'{x}'
            r[7].text = f'{x}'
            r[8].text = f'{x}'

        # Text sizes.
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        font = run.font
                        font.size = Pt(8)

    #-----------------------------------------------------------------------------------------------

    def add_temporary_team_table(self, w, cx):
        """
        Add table for temporary team.

        Parameters
        ----------
        w : Worksheet
            Worksheet.
        cx : ComplexTheme
            Complex theme.
        """

        # number of people
        n = len(w.lines)

        # table and its style
        t = self.doc.add_table(rows = n + 1, cols = 7)
        t.style = 'Table Grid'

        # head
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'ФИО'
        h[2].text = 'Должность'
        h[3].text = 'Табельный номер'
        h[4].text = 'Год рождения'
        h[5].text = 'Подразделение'
        h[6].text = 'Наименование подтем/тематик исследований'

        # add rest rows
        for i in range(n):
            r = t.rows[i + 1].cells
            wl = w.lines[i]
            e = wl.employee
            p = e.personal
            r[0].text = str(i + 1) + '.'
            r[1].text = p.surname_name_patronymic()
            r[2].text = wl.job_title.name
            r[3].text = e.tabel
            r[4].text = str(p.year)
            r[5].text = wl.job_place.half_full_name
            r[6].text = cx.all_thematics_titles()

    #-----------------------------------------------------------------------------------------------
    # Generate equipment methods.
    #-----------------------------------------------------------------------------------------------

    def add_equipment_title(self, cx):
        """
        Add title for equipment.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        """

        text = 'ПЕРЕЧЕНЬ\n'\
               'объектов особо ценного движимого имущества, используемого в процессе\n'\
               'выполнения научно-исследоввательской работы (основных средств и\n'\
               'нематериальных активов, амортизируемых в процессе выполнения работы)\n'\
               f'по комплексной теме {cx.title}'
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #-----------------------------------------------------------------------------------------------

    def add_inner_equipment_table(self, cx):
        """
        Add inner equipment table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        """

        # Add table title.
        self.add_paragraph('1. Перечень находящихся в оперативном управлении НИЦ «Курчатовский институт» объектов особо ценного имущества',
                           alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, is_bold=True)

        # Table and its style.
        t = self.doc.add_table(rows=4, cols=4)
        t.style = 'Table Grid'

        # All thematics.
        all_thematics_text = ''
        for th in cx.thematics:
            if all_thematics_text != '':
                all_thematics_text = all_thematics_text + '\n\n'
            all_thematics_text = all_thematics_text + f'Тематика\n{th.title}'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование объема особо ценного движимого имущества'
        h[2].text = 'Местоположение (здание, помещение)'
        h[3].text = 'Наименование подтем и тематик исследований'

        # CLK
        h = t.rows[1].cells
        h[0].text = '1.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П ОП2, '\
                    'подраздел МВС-10П ОП2 CLK (Cascade Lake),\n'\
                    f'{cx.clk_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # KNL
        h = t.rows[2].cells
        h[0].text = '2.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П МП2 (Knights Landing),\n'\
                    f'{cx.knl_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        # ICL
        h = t.rows[3].cells
        h[0].text = '3.'
        h[1].text = 'Суперкомпьютер МВС-10П ОП, раздел МВС-10П ОП3, '\
                    'подраздел МВС-10П ОП3 ICL (Ice Lake),\n'\
                    f'{cx.icl_need} узлочасов.'
        h[2].text = 'Москва, Ленинский проспект, 32А.'
        h[3].text = all_thematics_text

        set_table_columns_widths(t, [0.5, 2.0, 1.5, 3.5])

    #-----------------------------------------------------------------------------------------------

    def add_outer_equipment_table(self, cx):
        """
        Add outer equipment table.

        Parameters
        ----------
        cx : ComplexTheme
            Complex theme.
        """

        # Add table title.
        self.add_paragraph('2. Перечень объектов особо ценного движимого имущества, которое планируется привлечь на условиях аренды',
                           alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, is_bold=True)

        # Table and its style.
        t = self.doc.add_table(rows=2, cols=4)
        t.style = 'Table Grid'

        # Head.
        h = t.rows[0].cells
        h[0].text = '№ п/п'
        h[1].text = 'Наименование объема особо ценного движимого имущества'
        h[2].text = 'Местоположение (здание, помещение)'
        h[3].text = 'Наименование подтем и тематик исследований'

        set_table_columns_widths(t, [0.5, 2.0, 1.5, 3.5])

#===================================================================================================

def generate_order(cx, y, out):
    """
    Generate order.

    Parameters
    ----------
    cx : ComplexTheme
        Complex theme.
    y : int
        Year.
    out : str
        Outfile.
    """

    w = GeneratorWord()
    w.add_empty_lines(8)
    #
    w.add_paragraph('О проведении фундаментальной научно-исследовательской работы\n'
                    f'по комплексной теме {cx.title}', WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_empty_line()
    #
    p = w.add_paragraph('\tВ целях обеспечения реализации НИЦ «Курчатовский институт» '
                        'тематического плана научно-исследовательских и опытно-конструкторских '
                        f'работ в рамках выполнения государственного задания на {y} год и '
                        f'плановый период {y + 1} и {y + 2} годов (далее - '
                        'государственное задание), а также достижения показателей, '
                        'характеризующих качество работ по государственному заданию, ')
    r = p.add_run('п р и к а з ы в а ю:')
    r.bold = True
    #
    w.add_paragraph('\t1. Провести научно-исследовательскую работу по комплексной теме '
                    f'{cx.title} (далее - НИР).')
    #
    w.add_paragraph('\t2. Определить ответственным:')
    title_name = cx.deputy.job_title.name_r
    if title_name == 'руководителя отделения':
        title_name = title_name.split()[0]
    jobplace_name = cx.deputy.job_place.name_r
    pers = cx.deputy.employee.personal
    perss = pers.surname('ru')
    persn = pers.name_first_letter('ru')
    persp = pers.patronymic_first_letter('ru')
    w.add_paragraph('\t2.1. Должностным лицом за проведение НИР '
                    f'{title_name} {jobplace_name} {perss}а {persn}.{persp}.')
    #
    w.add_paragraph('\t2.2 Структурным подразделением Центра за выполнение НИР отделение '
                    'суперкомпьютерных систем и параллельных вычислений.')
    #
    title_name = cx.manager.job_title.name_r
    if title_name == 'руководителя отделения':
        title_name = title_name.split()[0]
    jobplace_name = cx.manager.job_place.name_r
    pers = cx.manager.employee.personal
    perss = pers.surname('ru')
    persn = pers.name_first_letter('ru')
    persp = pers.patronymic_first_letter('ru')
    w.add_paragraph(f'\t2.3. Руководителем НИР '
                    f'{title_name} {jobplace_name} {perss}а {persn}.{persp}.')
    #
    w.add_paragraph('\t3. Утвердить:')
    w.add_paragraph('\t3.1. Техническое задание на выполнение НИР (далее - ТЗ) согласно '
                    'приложению № 1 к настоящему приказу.')
    w.add_paragraph('\t3.2. Календарный план на выполнение НИР (далее - КП) согласно '
                    'приложению № 2 к настоящему приказу.')
    w.add_paragraph('\t3.3. Смету на выполнение НИР (далее - Смета) согласно '
                    'приложению № 3 к настоящему приказу.')
    w.add_paragraph('\t3.4. Состав временного трудового коллектива для выполнения НИР согласно '
                    'приложению № 4 к настоящему приказу.')
    w.add_paragraph('\t3.5. Перечень объектов особо ценного движимого имущества, используемого '
                    'в процессе выполнения НИР (основных средств и нематериальных активов, '
                    'амортизируемых в процессе выполнения работы) согласно '
                    'приложению № 5 к настоящему приказу.')
    w.add_paragraph(f'\t4. Установить срок выполнения НИР не позднее 31 декабря {y + 2} года. '
                    'Срок выполнения работ по настоящему приказу - '
                    f'не позднее 31 декабря {y} года.')
    w.add_paragraph('\t5. Ответственному должностному лицу за проведение НИР, указанному в '
                    'пункте 2.1 настоящего приказа, обеспечить:')
    w.add_paragraph('\tкоординацию выполнения НИР;')
    w.add_paragraph('\tконтроль соблюдения требований и показателей, характеризующих '
                    'качество работы, установленных ТЗ и КП;')
    w.add_paragraph('\tэффективное и целевое расходование средств, предусмотренных на '
                    'финансовое обеспечение государственного задания на выполнение работ в '
                    'пределах доведенных Центру лимитов бюджетных обязательств, предусмотренных '
                    'на выполнение НИР в соответствии со Сметой;')
    w.add_paragraph('\tподготовку единого отчета о НИР, оформленного в соответствии с '
                    'ГОСТ 7.32-2017 и проектов Форм направления сведений для размещения в '
                    'ЕГИСУ НИОКТР;')
    w.add_paragraph('\tполучение положительного заключения профильного экспертного совета '
                    'при Ученом совете НИЦ «Курчатовкий институт» на единый отчет о НИР;')
    w.add_paragraph(f'\tв срок до 15 января {y + 1} г. представление главному ученому секретарю '
                    'Центра единого отчета о НИР, оформленного в соответствии с ГОСТ 7.32-2017 '
                    'с положительным заключением экспертного совета при ученом совете Центра '
                    'и проектов Форм направления сведений для размещения в ЕГИСУ НИОКТР.')
    w.add_paragraph(f'\t6. Главному ученому секретарю Центра в срок до 1 февраля {y + 1} г. '
                    'обеспечить размещение отчета о НИР, Форм направления сведений в ЕГИСУ '
                    'НИОКТР и направление на экспертизу в соответствии с Правилами '
                    'осуществления федеральным государственным бюджетным учреждением '
                    '«Российская академия наук» научного и научно-методического руководства '
                    'научной и научно-технической деятельностью научных организаций и '
                    'образовательных организаций высшего образования, а также экспертизы '
                    'научных и научно-технических результатов, полученных этими организациями, '
                    'утвержденные постановлением Правительства Российской Федерации '
                    'от 30 декабря 2018 г. № 1781.')
    w.add_paragraph('\t7. Ответственным руководителям НИР, указанным в пункте 2.3 '
                    'настоящего приказа, обеспечить:')
    w.add_paragraph('\tорганизационное и научно-техническое руководство НИР;')
    w.add_paragraph('\tдостижение показателей, характеризующих качество работ в рамках '
                    'государственного задания, в соответствии с ТЗ и КП;')
    w.add_paragraph('\tсвоевременные подготовку и представление отчетной документации '
                    'в соответствии с ТЗ и КП для формирования единого отчета о НИР;')
    w.add_paragraph('\tобеспечение и подачу заявки на оформление в установленном порядке '
                    'государственной регистрации НИР и результатов интеллектуальной деятельности;')
    w.add_paragraph('\tвыполнение НИР в соответствии со Сметой.')
    w.add_paragraph('\t8. Контроль исполнения настоящего приказа оставляю за собой.')
    #
    w.add_empty_line()
    #
    w.add_signatures([wsl.dyakova_ya], False)
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_technical_task(n, cx, y, out):
    """
    Generate technical task document.

    Parameters
    ----------
    n : int
        Supplement number.
    cx : ComplexTheme
        Complex theme.
    y : int
        Start year.
    out : str
        Out file name.
    """

    w = GeneratorWord()
    #w.add_corner_inscription_supplement_to_order(n)
    #w.add_empty_line()
    w.add_technical_task_title(cx, y)
    w.add_empty_line()
    w.add_complex_theme_characteristics(cx, y)
    w.add_empty_line()
    w.add_signatures([cx.manager, wsl.shabanov_bm])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_calendar_plan(n, cx, y, out):
    """
    Generate calendar plan document.

    Parameters
    ----------
    n : int
        Supplement number.
    cx : ComplexTheme
        Complex theme.
    y : int
        Start year.
    out : str
        Out file name.
    """

    # Create outlays for thematics.
    for th in cx.thematics:
        th.outlay = outlay_tree.duplicate_outlay(cx.outlay, 'тематике исследований',
                                                 0.01 * th.funding_part(y))

    w = GeneratorWord()

    # Landscape.
    section = w.doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    #w.add_corner_inscription_supplement_to_order(n)
    #w.add_empty_line()
    w.add_calendar_plan_title(cx, y)
    w.add_empty_line()
    w.add_calendar_plan_table(cx, y)
    w.add_empty_line()
    w.add_calendar_plan_table(cx, y + 1)
    w.add_empty_line()
    w.add_calendar_plan_table(cx, y + 2)
    w.add_empty_line()
    w.add_signatures([cx.manager, wsl.shabanov_bm])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_form_gos_assignment_suppl_09_pre_outlay(cx, y, out):
    """
    Generate
    'form gos assignment, supplement 09 - pre outlay'.

    Parameters
    ----------
    cx : ComplexTheme
        Complex theme.
    y : int
        Year.
    out : str
        Out file.
    """

    # Create document.
    w = GeneratorWord()

    # Title.
    w.add_paragraph('ПРЕДВАРИТЕЛЬНАЯ СМЕТА\n '
                    'прямых и общепроизводственных расходов, непосредственно связанных с '
                    'выполнением научно-исследовательской работы '
                    f'по комплексной теме {cx.title} на очередной {y} год '
                    f'и плановый период {y + 1} и {y + 2} годов',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)

    # First table.
    w.add_paragraph('1. Предварительная смета прямых и общепроизводственных расходов, '
                    'непосредственно связанных с выполнением научно-исследовательской работы '
                    f'на очередной {y} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_outlay_table_theme_with_subthemes(cx.outlay, y, False)
    w.add_empty_line()

    # Second table.
    w.add_paragraph('2. Предварительная смета прямых и общепроизводственных расходов, '
                    'непосредственно связанных с выполнением научно-исследовательской работы '
                    f'на первый плановый год - {y + 1} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_outlay_table_theme_with_subthemes(cx.outlay, y + 1, False)
    w.add_empty_line()

    # Third table.
    w.add_paragraph('3. Предварительная смета прямых и общепроизводственных расходов, '
                    'непосредственно связанных с выполнением научно-исследовательской работы '
                    f'на второй плановый год - {y + 2} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_outlay_table_theme_with_subthemes(cx.outlay, y + 2, False)
    w.add_empty_line()

    # Add signatures and close file.
    w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_outlay(n, cx, y, out):
    """
    Generate outlay document.

    Parameters
    ----------
    n : int
        Supplement number.
    cx : ComplexTheme
        Complex theme.
    y : int
        Start year.
    out : str
        Out file name.
    """

    w = GeneratorWord()
    #w.add_corner_inscription_supplement_to_order(n)
    #w.add_empty_line()
    w.add_outlay_title(cx, y)
    w.add_empty_line()

    # Add short outlay for complex theme.
    w.add_paragraph(f'1. ПРЕДВАРИТЕЛЬНАЯ СМЕТА\nпо комплексной теме {cx.title}',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_short_outlay_table(cx, y)
    w.add_empty_line()

    # Add outlay for complex theme.
    w.add_paragraph('1. ПРЕДВАРИТЕЛЬНАЯ СМЕТА\n'
                    f'на очередной {y} год и плановый период {y + 1} и {y + 2} годов\n'
                    f'по комплексной теме {cx.title}',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_outlay_table(cx.outlay, y)
    w.add_empty_line()

    # Add outlays for thematic.
    for i, th in enumerate(cx.thematics):
        w.add_paragraph(f'1.{i + 1}. ПРЕДВАРИТЕЛЬНАЯ СМЕТА\n'
                        f'на очередной {y} год и плановый период {y + 1} и {y + 2} годов\n'
                        f'по тематике исследований {th.title}',
                        WD_PARAGRAPH_ALIGNMENT.CENTER, True)
        w.add_outlay_table(th.outlay, y)
        w.add_empty_line()

    w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_form_gos_assignment_suppl_10_temporary_team(cx, team, y, out):
    """
    Generate form for gos assignment.
    Supplement 10 (temporary team).

    Parameters
    ----------
    cx : ComplexTheme
        Complex theme.
    team : Worksheet
        Worksheet.
    y : int
        Year.
    out : str
        Out file.
    """

    # Create document.
    w = GeneratorWord()

    # Title.
    w.add_paragraph('СОСТАВ\n'
                    'временного трудового коллектива для выполнения научно-исследовательской '
                    f'работы и планируемые трудозатраты на очередной {y} год '
                    f'и плановый период {y + 1} и {y + 2} годов по комплексной теме {cx.title}',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)

    # First year.
    w.add_paragraph('1. Состав временного трудового коллектива и планируемые трудозатраты '
                    f'на первый год планового периода - {y} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_temporary_team_table_with_workload(team, cx)
    w.add_empty_line()

    # First year.
    w.add_paragraph('2. Состав временного трудового коллектива и планируемые трудозатраты '
                    f'на первый год планового периода - {y + 1} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_temporary_team_table_with_workload(team, cx)
    w.add_empty_line()

    # First year.
    w.add_paragraph('3. Состав временного трудового коллектива и планируемые трудозатраты '
                    f'на первый год планового периода - {y + 2} год',
                    WD_PARAGRAPH_ALIGNMENT.CENTER, True)
    w.add_temporary_team_table_with_workload(team, cx)
    w.add_empty_line()

    # Add signatures and close file.
    w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_temporary_team(n, cx, team, y, out):
    """
    Generate temporary team document.

    Parameters
    ----------
    n : int
        Supplement number.
    cx : ComplexTheme
        Complex theme.
    team : Worksheet
        Team working on this theme.
    y : int
        Start year.
    out : str
        Out file name.
    """

    w = GeneratorWord()
    #w.add_corner_inscription_supplement_to_order(n)
    #w.add_empty_line()
    w.add_temporary_team_title(cx, y)
    w.add_empty_line()
    w.add_temporary_team_table(team, cx)
    w.add_empty_line()
    w.add_signatures([cx.manager, wsl.shabanov_bm, wsl.smirnnova_oe, wsl.petrischev_av])
    w.save(out + '.docx')

#---------------------------------------------------------------------------------------------------

def generate_equipment(n, cx, out):
    """
    Generate equipment document.

    Parameters
    ----------
    n : int
        Supplement number.
    cx : ComplexTheme
        Complex theme.
    out : str
        Out file name.
    """

    w = GeneratorWord()
    #w.add_corner_inscription_supplement_to_order(n)
    #w.add_empty_line()
    w.add_equipment_title(cx)
    w.add_empty_line()
    w.add_inner_equipment_table(cx)
    w.add_empty_line()
    w.add_outer_equipment_table(cx)
    w.add_empty_line()
    w.add_signatures([cx.manager, wsl.shabanov_bm])
    w.save(out + '.docx')

#===================================================================================================

if __name__ == '__main__':
    pass

#===================================================================================================
